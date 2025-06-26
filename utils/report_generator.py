import os
from dotenv import load_dotenv
load_dotenv()

from docx import Document
from job_matching.filters import passes_hard_filters

def generate_cv_report(cv_id: str, cv_data: dict, offer_id: str, offer_text: str, criteria: dict, output_dir: str):
    doc = Document()
    doc.add_heading(f"Report: CV {cv_id} vs Offer {offer_id}", level=1)
    doc.add_paragraph(f"Offer:\n{offer_text[:500]}...")

    doc.add_heading("CV Structured Data", level=2)
    for k, v in cv_data.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("Hard Filters", level=2)
    passed = passes_hard_filters(cv_data, criteria)
    status = "PASSED" if passed else "FAILED"
    doc.add_paragraph(f"Overall hard filters: {status}")

    if not passed:
        doc.add_heading("Missing Requirements", level=3)
        for skill in criteria.get("required_skills", []):
            if skill.lower() not in cv_data.get("skills", "").lower():
                doc.add_paragraph(f"- Missing skill: {skill}")

    out_path = os.path.join(output_dir, f"report_{cv_id}_{offer_id}.docx")
    os.makedirs(output_dir, exist_ok=True)
    doc.save(out_path)
    return out_path
