import os
from docx import Document
from langchain_openai import ChatOpenAI
from langchain.schema import HumanMessage
from dotenv import load_dotenv

load_dotenv()


# instantiate LLM once
llm = ChatOpenAI(model="gpt-3.5-turbo")

def generate_company_cv(data: dict,
                        template_path: str,
                        output_path: str,
                        match_info: dict = None) -> str:
    """
    Uses GPT to fill a company CV template with parsed candidate data
    + optional match_info. Supports .docx and .txt templates.
    Returns the full path to the generated file.
    """
    # ensure output dir
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    base, ext = os.path.splitext(template_path)
    ext = ext.lower()

    # 1) Read template into a text block
    if ext == ".docx":
        tpl_doc = Document(template_path)
        tpl_lines = [p.text for p in tpl_doc.paragraphs]
        template_str = "\n".join(tpl_lines)
    elif ext == ".txt":
        with open(template_path, "r", encoding="utf-8") as f:
            template_str = f.read()
    else:
        raise ValueError(f"Unsupported template format: {ext}")

    # 2) Build prompt for GPT
    prompt = (
        "You are a CV formatting assistant. "
        "Fill the following company CV template by replacing placeholders with the candidate's data.\n\n"
        f"Template:\n{template_str}\n\n"
        f"Candidate data (JSON):\n{data}\n\n"
    )
    if match_info:
        prompt += f"Match info:\n{match_info}\n\n"
    prompt += (
        "Preserve the template structure and only replace placeholders like {{field}}. "
        "Output the fully filled CV text."
    )

    # 3) Call GPT
    response = llm.invoke([HumanMessage(content=prompt)])
    filled = response.content

    # 4) Write out the result
    if ext == ".docx":
        out_doc = Document()
        for line in filled.split("\n"):
            out_doc.add_paragraph(line)
        out_doc.save(output_path)
    else:  # .txt
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(filled)

    return output_path
